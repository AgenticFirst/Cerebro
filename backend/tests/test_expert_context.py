"""Per-expert context-files CRUD."""

from __future__ import annotations

import os

import pytest

from main import app


@pytest.fixture()
def ctx_client(client, tmp_path):
    files_dir = str(tmp_path / "files")
    parsed_dir = os.path.join(files_dir, "_parsed")
    os.makedirs(parsed_dir, exist_ok=True)
    app.state.files_dir = files_dir
    app.state.parsed_files_dir = parsed_dir
    return client


def _mk_expert(c, name="Manual Author") -> str:
    res = c.post("/experts", json={"name": name, "description": "writes manuals"})
    assert res.status_code == 201, res.text
    return res.json()["id"]


def _mk_docx(tmp_path, name="estructura.docx") -> str:
    from docx import Document

    doc = Document()
    doc.add_heading("Estructura", level=1)
    doc.add_paragraph("guía maestra para los manuales")
    p = str(tmp_path / name)
    doc.save(p)
    return p


def _register_file(c, file_path: str) -> str:
    res = c.post(
        "/files/items/from-path",
        json={"file_path": file_path, "source": "expert-context"},
    )
    assert res.status_code == 201, res.text
    return res.json()["id"]


def test_attach_list_detach_context_file(ctx_client, tmp_path):
    eid = _mk_expert(ctx_client)
    fp = _mk_docx(tmp_path)
    fid = _register_file(ctx_client, fp)
    # Pre-parse so char_count is recorded.
    parsed = ctx_client.post("/files/parse", json={"file_path": fp}).json()
    assert parsed["char_count"] > 0

    # Attach
    res = ctx_client.post(
        f"/experts/{eid}/context-files",
        json={"file_item_id": fid, "kind": "template"},
    )
    assert res.status_code == 201, res.text
    body = res.json()
    assert body["expert_id"] == eid
    assert body["file_item_id"] == fid
    assert body["kind"] == "template"
    assert body["file_name"] == "estructura.docx"
    assert body["char_count"] == parsed["char_count"]
    assert body["truncated"] is False
    ctx_id = body["id"]

    # List
    listed = ctx_client.get(f"/experts/{eid}/context-files")
    assert listed.status_code == 200
    rows = listed.json()
    assert len(rows) == 1
    assert rows[0]["id"] == ctx_id
    assert rows[0]["file_name"] == "estructura.docx"

    # Patch
    patched = ctx_client.patch(
        f"/experts/{eid}/context-files/{ctx_id}",
        json={"kind": "reference", "sort_order": 5.5},
    )
    assert patched.status_code == 200
    assert patched.json()["kind"] == "reference"
    assert patched.json()["sort_order"] == 5.5

    # Detach
    deleted = ctx_client.delete(f"/experts/{eid}/context-files/{ctx_id}")
    assert deleted.status_code == 204
    final = ctx_client.get(f"/experts/{eid}/context-files")
    assert final.json() == []


def test_sort_order_defaults_to_max_plus_one(ctx_client, tmp_path):
    eid = _mk_expert(ctx_client)
    a = _register_file(ctx_client, _mk_docx(tmp_path, name="a.docx"))
    b = _register_file(ctx_client, _mk_docx(tmp_path, name="b.docx"))
    c = _register_file(ctx_client, _mk_docx(tmp_path, name="c.docx"))

    r1 = ctx_client.post(
        f"/experts/{eid}/context-files", json={"file_item_id": a}
    ).json()
    r2 = ctx_client.post(
        f"/experts/{eid}/context-files", json={"file_item_id": b}
    ).json()
    r3 = ctx_client.post(
        f"/experts/{eid}/context-files", json={"file_item_id": c}
    ).json()

    assert r1["sort_order"] == 0.0
    assert r2["sort_order"] == 1.0
    assert r3["sort_order"] == 2.0


def test_attach_rejects_unknown_expert(ctx_client, tmp_path):
    fid = _register_file(ctx_client, _mk_docx(tmp_path))
    res = ctx_client.post(
        "/experts/nope/context-files", json={"file_item_id": fid}
    )
    assert res.status_code == 404


def test_attach_rejects_unknown_file(ctx_client):
    eid = _mk_expert(ctx_client)
    res = ctx_client.post(
        f"/experts/{eid}/context-files", json={"file_item_id": "missing"}
    )
    assert res.status_code == 404


def test_attach_rejects_invalid_kind(ctx_client, tmp_path):
    eid = _mk_expert(ctx_client)
    fid = _register_file(ctx_client, _mk_docx(tmp_path))
    res = ctx_client.post(
        f"/experts/{eid}/context-files",
        json={"file_item_id": fid, "kind": "bogus"},
    )
    assert res.status_code == 400


def test_list_skips_soft_deleted_files(ctx_client, tmp_path):
    eid = _mk_expert(ctx_client)
    fid = _register_file(ctx_client, _mk_docx(tmp_path))
    ctx_client.post(
        f"/experts/{eid}/context-files", json={"file_item_id": fid}
    )
    ctx_client.delete(f"/files/items/{fid}")
    rows = ctx_client.get(f"/experts/{eid}/context-files").json()
    assert rows == []
